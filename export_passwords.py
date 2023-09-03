from hashlib import md5
from docx import Document
from docx.shared import Cm
import requests as rq
import argparse
import dotenv
import os

dotenv.load_dotenv(".env")
parser = argparse.ArgumentParser(
    prog="python3 export_passwords.py",
    description="Export default passwords using hash of logins and sends word file using telegram bot",
)

parser.add_argument("-f", "--filename", default="passwords.docx")
parser.add_argument("-u", "--userid", default=os.environ.get("EXPORT_USER_ID"))
parser.add_argument("-b", "--bottoken", default=os.environ.get("EXPORT_BOT_TOKEN"))
parser.add_argument("-s", "--saveonly", action="store_true")
parser.add_argument("-t", "--teachers", action="store_true")

args = parser.parse_args()


def get_default_password(login):
    return md5(login.encode()).hexdigest()[:6]


from backendtypes import DataBase

with DataBase() as db:
    if args.teachers:
        classes=["Классные руководители"]
        db.q.execute("SELECT login FROM users WHERE is_teacher = 1")
        logins_bcls=[[lgn for lgn, in db.q.fetchall()]]
    else:
        classes = db.get_all_classes()
        logins_bcls = []
        for classr in classes:
            users = db.get_users_by_class(classr)
            logins_bcls.append([user.login for user in users])

doc = Document()

j = len(classes) - 1

for ind, (classr, logins) in enumerate(zip(classes, logins_bcls)):
    doc.add_heading(classr, 1)
    doc.add_paragraph("diary130.ru")
    table = doc.add_table(rows=1, cols=2)
    row = table.rows[0].cells
    for i, item in enumerate(["Логин", "Пароль"]):
        p = row[i].paragraphs[0]
        p.add_run(item).bold = True

    for login in logins:
        row = table.add_row().cells
        row[0].text = login
        row[1].text = get_default_password(login)
    table.style = "Table Grid"
    if ind < j:
        doc.add_page_break()

sections = doc.sections
for section in sections:
    section.top_margin = Cm(1)
    section.bottom_margin = Cm(1)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)
doc.save(args.filename)

if not (args.saveonly):
    rq.post(
        f"https://api.telegram.org/bot{args.bottoken}/sendDocument",
        data={"chat_id": args.userid},
        files={"document": open(args.filename, "rb")},
    )
